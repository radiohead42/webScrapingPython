import bs4
import requests
from docx import Document
from PIL import Image

doc = Document()

url_base = 'https://es.wikipedia.org/wiki/Intel'

baseHTML = bs4.BeautifulSoup(requests.get(url_base).text, 'lxml')

titulo = baseHTML.find('h1').text

nombreImagen = titulo + '.jpg'

primerParrafo = baseHTML.select('p')[0].text

primerSubtitulo = baseHTML.select('h2')[1].text

contenido = baseHTML.select('p')[1].text

final = baseHTML.select('p')[2].text

imagen = baseHTML.select('figure a img')[0]['src']

imagen = 'https:' + imagen

textoImagen = baseHTML.select('figure a img')[0]

cadenaCompleta = titulo + '\n' + primerParrafo + '\n' + primerSubtitulo + '\n' + contenido

imagenes_curso = requests.get(imagen)
f = open(nombreImagen,'wb')
f.write(imagenes_curso.content)
f.close()

def convertir_jpg_a_jpeg(archivo_entrada, archivo_salida):
    try:
        # Abrir la imagen en formato jpg
        imagen = Image.open(archivo_entrada)

        # Guardar la imagen en formato jpeg
        imagen.save(archivo_salida, "JPEG")

        print("La imagen se ha convertido exitosamente a formato JPEG.")
    except Exception as e:
        print("Ocurrió un error al convertir la imagen:", e)


# Ejemplo de uso
archivo_entrada = nombreImagen
archivo_salida = titulo + ".jpeg"
convertir_jpg_a_jpeg(archivo_entrada, archivo_salida)

print(nombreImagen)

doc.add_heading(titulo, level=1)

doc.add_paragraph(primerParrafo)

doc.add_heading(primerSubtitulo, level=2)

doc.add_paragraph(contenido)

doc.add_paragraph(final)

doc.add_picture(titulo+'.jpeg')

doc.add_paragraph(textoImagen)

doc.save('tarea.docx')

#print(cadenaCompleta)

